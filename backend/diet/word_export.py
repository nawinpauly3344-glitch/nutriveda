"""
NutriVeda Diet Plan — Word Document Generator
Generates beautiful, editable .docx files with professional styling.
Inspired by HealthifyMe, MyFitnessPal, and clinical nutrition report design.
"""

import os
import re
import logging
from pathlib import Path
from datetime import datetime
from typing import List, Tuple, Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

log = logging.getLogger(__name__)

# ── Output directory ──────────────────────────────────────────────────────────
WORD_DIR = Path(__file__).parent.parent / "word_plans"
WORD_DIR.mkdir(exist_ok=True)

# ── Professional Color Palette ────────────────────────────────────────────────
# All table column headers use C_HDR_BG (very dark navy) + C_WHITE text.
# This guarantees maximum readability regardless of section.

# Table header — universal dark navy for all tables
C_HDR_BG    = RGBColor(0x1E, 0x3A, 0x5F)   # Deep navy — column headers (always visible)

# Primary brand — deep emerald green
C_PRIMARY   = RGBColor(0x06, 0x5F, 0x46)   # Dark emerald (section banners, cover)
C_DARK_GRN  = RGBColor(0x02, 0x2C, 0x22)   # Almost-black green (cover accent)
C_MID_GRN   = RGBColor(0x05, 0x96, 0x69)   # Emerald mid (borders, accents)
C_LIGHT_GRN = RGBColor(0xEC, 0xFD, 0xF5)   # Very pale mint (section bg tint)
C_PALE_GRN  = RGBColor(0xF0, 0xFF, 0xF8)   # Near-white mint (alternating rows)

# Section heading banner backgrounds (DARK — white text always readable)
C_SEC_GREEN  = RGBColor(0x06, 0x5F, 0x46)   # Nutrition / meals
C_SEC_NAVY   = RGBColor(0x1E, 0x3A, 0x5F)   # Exercise / training
C_SEC_PURPLE = RGBColor(0x3B, 0x0F, 0x6B)   # Supplements / vitamins
C_SEC_TEAL   = RGBColor(0x0A, 0x40, 0x50)   # Hydration
C_SEC_AMBER  = RGBColor(0x6B, 0x21, 0x0A)   # Cheat meals / indulgence
C_SEC_SLATE  = RGBColor(0x1E, 0x29, 0x3B)   # Tips / habits / general
C_SEC_RED    = RGBColor(0x7F, 0x1D, 0x1D)   # Avoid / restrict

# Meal row tints (very light — do NOT clash with dark header)
C_ROW_BF    = RGBColor(0xFF, 0xFB, 0xEB)   # Breakfast — warm amber tint
C_ROW_LN    = RGBColor(0xFF, 0xF7, 0xED)   # Lunch — warm peach tint
C_ROW_DN    = RGBColor(0xEF, 0xF6, 0xFF)   # Dinner — cool blue tint
C_ROW_SN    = RGBColor(0xF0, 0xFF, 0xF4)   # Snack / mid — mint tint
C_ROW_TOT   = RGBColor(0xF1, 0xF5, 0xF9)   # Daily total — neutral slate tint
C_ROW_ALT   = RGBColor(0xFA, 0xFA, 0xFA)   # Alternating row (near-white)

# Tip / blockquote box
C_TIP_BG    = RGBColor(0xFF, 0xFB, 0xEB)   # Warm amber background
C_TIP_TEXT  = RGBColor(0x6B, 0x21, 0x0A)   # Deep amber text

# Neutrals
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_TEXT      = RGBColor(0x1A, 0x20, 0x2C)   # Near-black (slightly warm)
C_TEXT_GRAY = RGBColor(0x64, 0x74, 0x8B)   # Cool gray
C_BORDER    = "CBD5E1"                      # Soft slate border


# ── XML Helpers ───────────────────────────────────────────────────────────────

def _set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(color))
    # Remove existing shd if present
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(shd)


def _set_table_borders(table, color: str = C_BORDER, size: int = 4):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(borders)


def _cell_text(cell, text: str, bold=False, size=10, color=None,
               align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Pt(4)
    p.paragraph_format.right_indent = Pt(4)
    run = p.add_run(str(text) if text else "")
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color or C_TEXT


def _add_colored_heading(doc: Document, text: str, level: int,
                          banner_bg: Optional[RGBColor] = None,
                          icon: str = ""):
    """Add a section heading.
    h1/h2: full-width dark banner with white text.
    h3: bold text in section color with a thin left accent line via paragraph border.
    """
    sizes = {1: 16, 2: 13, 3: 11, 4: 10}

    if level <= 2 and banner_bg is not None:
        # Full-width DARK banner — white text, always readable
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(table, color=str(banner_bg), size=0)
        cell = table.cell(0, 0)
        _set_cell_bg(cell, banner_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        label = f"  {text}" if not icon else f"  {text}"
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(sizes.get(level, 12))
        run.font.color.rgb = C_WHITE
        _add_spacer(doc, 4)
    else:
        # h3/h4: simple bold text with section color
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        label = text
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(sizes.get(level, 10))
        run.font.color.rgb = banner_bg if banner_bg else C_TEXT


def _add_spacer(doc: Document, pts: int = 4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)


# ── Markdown Parser ───────────────────────────────────────────────────────────

def _clean_inline(text: str) -> str:
    """Remove markdown inline formatting."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"~~(.+?)~~", r"\1", text)
    return text.strip()


def _parse_bold_runs(text: str) -> List[Tuple[str, bool]]:
    """Return list of (text, is_bold) tuples for inline bold."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    result = []
    for p in parts:
        if p.startswith("**") and p.endswith("**"):
            result.append((p[2:-2], True))
        elif p:
            result.append((p, False))
    return result


def _parse_blocks(plan_text: str) -> List[dict]:
    """Parse plan markdown into structured blocks."""
    lines = plan_text.splitlines()
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        raw = line.strip()

        # Heading
        if raw.startswith("### "):
            blocks.append({"type": "h3", "text": raw[4:].strip()})
        elif raw.startswith("## "):
            blocks.append({"type": "h2", "text": raw[3:].strip()})
        elif raw.startswith("# "):
            blocks.append({"type": "h1", "text": raw[2:].strip()})

        # Table: collect consecutive | lines
        elif raw.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            # Parse rows, skip separator rows (|---|)
            rows = []
            for tl in table_lines:
                if re.match(r"^\|[-| :]+\|$", tl):
                    continue
                cells = [c.strip() for c in tl.split("|")[1:-1]]
                if cells:
                    rows.append(cells)
            if rows:
                blocks.append({"type": "table", "rows": rows})
            continue

        # Horizontal rule
        elif raw in ("---", "***", "___") or re.match(r"^-{3,}$", raw):
            blocks.append({"type": "divider"})

        # Bullet / numbered list
        elif raw.startswith("- ") or raw.startswith("• ") or re.match(r"^\d+\.", raw):
            text = re.sub(r"^[-•]\s+", "", raw)
            text = re.sub(r"^\d+\.\s+", "", text)
            blocks.append({"type": "bullet", "text": text})

        # Blockquote / tip
        elif raw.startswith(">"):
            text = raw.lstrip("> ").strip()
            blocks.append({"type": "tip", "text": text})

        # Empty
        elif not raw:
            blocks.append({"type": "spacer"})

        # Normal paragraph
        else:
            blocks.append({"type": "para", "text": raw})

        i += 1
    return blocks


# ── Section color resolver ────────────────────────────────────────────────────

def _section_colors(heading_text: str):
    """Return (banner_bg, icon) — all banners use dark bg + white text."""
    t = heading_text.lower()
    if any(k in t for k in ["exercise", "workout", "fitness", "training", "sport"]):
        return C_SEC_NAVY, "exercise"
    if any(k in t for k in ["supplement", "vitamin", "mineral", "probiotic"]):
        return C_SEC_PURPLE, "supplement"
    if any(k in t for k in ["hydration", "water", "drink", "fluid"]):
        return C_SEC_TEAL, "hydration"
    if any(k in t for k in ["avoid", "restrict", "limit", "do not", "never"]):
        return C_SEC_RED, "avoid"
    if any(k in t for k in ["cheat", "indulg", "treat", "sunday"]):
        return C_SEC_AMBER, "cheat"
    if any(k in t for k in ["tip", "habit", "sleep", "recovery", "rest"]):
        return C_SEC_SLATE, "tip"
    if any(k in t for k in ["progress", "tracker", "measurement", "checkpoint"]):
        return C_SEC_TEAL, "progress"
    if any(k in t for k in ["message", "personal", "note", "nutritionist"]):
        return C_SEC_GREEN, "message"
    if any(k in t for k in ["blueprint", "nutritional", "macro", "calorie"]):
        return C_SEC_SLATE, "blueprint"
    # Default: emerald green (meal plan / general)
    return C_SEC_GREEN, "meal"


_SECTION_ICONS = {
    "exercise": "🏋",
    "supplement": "💊",
    "hydration": "💧",
    "avoid": "🚫",
    "cheat": "🎉",
    "tip": "💡",
    "progress": "📈",
    "message": "💬",
    "blueprint": "📊",
    "meal": "🍽",
}


def _is_meal_table(first_row: List[str]) -> bool:
    cols = [c.lower() for c in first_row]
    return any(k in " ".join(cols) for k in ["meal", "time", "food", "kcal", "day"])


def _meal_row_color(meal_label: str) -> Optional[RGBColor]:
    m = meal_label.lower()
    if "breakfast" in m or "morning" in m:
        return C_ROW_BF
    if "lunch" in m:
        return C_ROW_LN
    if "dinner" in m or "evening" in m:
        return C_ROW_DN
    if "snack" in m or "mid" in m:
        return C_ROW_SN
    if "total" in m or "daily" in m or "day" == m.strip():
        return C_ROW_TOT
    return None


# ── Cover Page ────────────────────────────────────────────────────────────────

def _build_cover(doc: Document, client_name: str, submission_id: int):
    # Dark emerald banner — brand header
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table, color=str(C_PRIMARY), size=0)
    cell = table.cell(0, 0)
    _set_cell_bg(cell, C_PRIMARY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("NutriVeda")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = C_WHITE

    p2 = cell.add_paragraph("Certified Nutrition Consultation")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(24)
    r2 = p2.runs[0]
    r2.font.size = Pt(13)
    r2.font.color.rgb = C_LIGHT_GRN
    r2.italic = True

    _add_spacer(doc, 20)

    # Subtitle line
    p3 = doc.add_paragraph("YOUR PERSONALISED 30-DAY TRANSFORMATION PLAN")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.runs[0]
    r3.bold = True
    r3.font.size = Pt(14)
    r3.font.color.rgb = C_PRIMARY

    _add_spacer(doc, 12)

    # Client info card — 2-column table
    table2 = doc.add_table(rows=3, cols=2)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table2, color=str(C_MID_GRN), size=6)

    data = [
        ("Prepared for", client_name),
        ("Date", datetime.now().strftime("%d %B %Y")),
        ("Reference", f"NV-{submission_id:04d}"),
    ]
    for row_idx, (label, value) in enumerate(data):
        row = table2.rows[row_idx]
        _set_cell_bg(row.cells[0], C_HDR_BG)
        _cell_text(row.cells[0], label, bold=True, size=10,
                   color=C_WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_bg(row.cells[1], C_LIGHT_GRN)
        _cell_text(row.cells[1], value, bold=(row_idx == 0), size=11,
                   color=C_TEXT, align=WD_ALIGN_PARAGRAPH.LEFT)

    _add_spacer(doc, 20)

    # Disclaimer box
    table3 = doc.add_table(rows=1, cols=1)
    _set_table_borders(table3, color=str(C_MID_GRN), size=4)
    cell3 = table3.cell(0, 0)
    _set_cell_bg(cell3, C_LIGHT_GRN)
    p4 = cell3.paragraphs[0]
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(8)
    p4.paragraph_format.space_after = Pt(8)
    r4 = p4.add_run(
        "This plan is prepared exclusively for the client named above. "
        "Please consult your nutritionist before making any modifications."
    )
    r4.font.size = Pt(9)
    r4.italic = True
    r4.font.color.rgb = C_PRIMARY

    doc.add_page_break()


# ── Table Renderer ────────────────────────────────────────────────────────────

def _render_table(doc: Document, rows: List[List[str]],
                  is_meal: bool = False):
    """Render a markdown table into the Word document.
    ALL table headers use C_HDR_BG (deep navy) + white text — always readable.
    """
    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    _set_table_borders(table, color=C_BORDER, size=4)

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        is_header = r_idx == 0
        is_total = any("total" in str(c).lower() or "daily total" in str(c).lower()
                       for c in row_data)

        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= n_cols:
                break
            cell = row.cells[c_idx]
            text = _clean_inline(str(cell_text))

            if is_header:
                # Universal dark navy header — maximum contrast, always visible
                _set_cell_bg(cell, C_HDR_BG)
                _cell_text(cell, text, bold=True, size=10,
                           color=C_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
            elif is_total:
                _set_cell_bg(cell, C_ROW_TOT)
                _cell_text(cell, text, bold=True, size=10, color=C_TEXT)
            elif is_meal:
                meal_label = _clean_inline(str(row_data[0]))
                bg = _meal_row_color(meal_label)
                if bg:
                    _set_cell_bg(cell, bg)
                elif r_idx % 2 == 0:
                    _set_cell_bg(cell, C_ROW_ALT)
                _cell_text(cell, text, size=10, color=C_TEXT,
                           bold=bool(c_idx == 0 and meal_label))
            else:
                if r_idx % 2 == 1:
                    _set_cell_bg(cell, C_ROW_ALT)
                _cell_text(cell, text, size=10, color=C_TEXT)

    _add_spacer(doc, 8)


# ── Main Document Builder ─────────────────────────────────────────────────────

def _render_blocks(doc: Document, blocks: List[dict]):
    """Render all parsed blocks into the Word document."""
    current_banner_bg = C_SEC_GREEN
    pending_table: Optional[List[List[str]]] = None

    def flush_table():
        nonlocal pending_table
        if pending_table:
            is_meal = _is_meal_table(pending_table[0])
            _render_table(doc, pending_table, is_meal=is_meal)
            pending_table = None

    _strip_emoji = lambda t: re.sub(r"^[\U0001F300-\U0001FFFF\u2600-\u27BF\u2700-\u27BF]+\s*", "", t).strip()

    for block in blocks:
        btype = block["type"]

        if btype == "table":
            if pending_table is None:
                pending_table = block["rows"]
            else:
                pending_table.extend(block["rows"])
            continue
        else:
            flush_table()

        if btype == "h1":
            text = block["text"]
            clean = re.sub(r"^[^\w]+", "", text).strip()
            banner_bg, _ = _section_colors(clean)
            current_banner_bg = banner_bg
            display = _strip_emoji(text)
            _add_colored_heading(doc, display, 1, banner_bg=banner_bg)

        elif btype == "h2":
            text = block["text"]
            clean = re.sub(r"^[^\w]+", "", text).strip()
            banner_bg, _ = _section_colors(clean)
            current_banner_bg = banner_bg
            display = _strip_emoji(text)
            _add_colored_heading(doc, display, 2, banner_bg=banner_bg)

        elif btype == "h3":
            text = block["text"]
            display = _strip_emoji(text)
            _add_colored_heading(doc, display, 3, banner_bg=current_banner_bg)

        elif btype == "divider":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), C_BORDER)
            pBdr.append(bottom)
            pPr.append(pBdr)

        elif btype == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Pt(8)
            for frag, is_bold in _parse_bold_runs(block["text"]):
                r = p.add_run(_clean_inline(frag) if not is_bold else frag)
                r.bold = is_bold
                r.font.size = Pt(10)
                r.font.color.rgb = C_TEXT

        elif btype == "tip":
            # Highlighted tip / blockquote box
            table = doc.add_table(rows=1, cols=1)
            _set_table_borders(table, color=str(C_SEC_AMBER), size=4)
            cell = table.cell(0, 0)
            _set_cell_bg(cell, C_TIP_BG)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Pt(6)
            runs = _parse_bold_runs(block["text"])
            for frag, is_bold in runs:
                r = p.add_run(_clean_inline(frag) if not is_bold else frag)
                r.bold = is_bold
                r.font.size = Pt(10)
                r.italic = not is_bold
                r.font.color.rgb = C_TIP_TEXT
            _add_spacer(doc, 6)

        elif btype == "para":
            raw = block["text"]
            if re.match(r"^[-=]{3,}$", raw.strip()):
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            for frag, is_bold in _parse_bold_runs(raw):
                r = p.add_run(_clean_inline(frag) if not is_bold else frag)
                r.bold = is_bold
                r.font.size = Pt(10)
                r.font.color.rgb = C_TEXT

        elif btype == "spacer":
            _add_spacer(doc, 4)

    flush_table()


# ── Page Setup ────────────────────────────────────────────────────────────────

def _configure_document(doc: Document):
    """Set margins, default font, page size."""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Default font
    from docx.oxml.ns import qn as _qn
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = C_TEXT


def _add_footer(doc: Document):
    """Add footer to all pages."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(
        "NutriVeda — Certified Nutrition Consultation  |  "
        "This plan is confidential and prepared exclusively for the named client."
    )
    run.font.size = Pt(7.5)
    run.font.color.rgb = C_TEXT_GRAY
    run.italic = True

    # Top border on footer
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), str(C_MID_GRN))
    pBdr.append(top)
    pPr.append(pBdr)


# ── Public API ────────────────────────────────────────────────────────────────

def generate_word_doc(
    plan_text: str,
    client_name: str,
    submission_id: int,
    plan_id: int,
) -> Optional[str]:
    """
    Generate a beautiful Word document from the plan markdown text.
    Returns the saved file path, or None on failure.
    """
    try:
        doc = Document()
        _configure_document(doc)
        _add_footer(doc)

        # Cover page
        _build_cover(doc, client_name, submission_id)

        # Parse and render plan content
        blocks = _parse_blocks(plan_text)
        _render_blocks(doc, blocks)

        # Save
        safe_name = re.sub(r"[^\w]", "_", client_name.lower())[:30]
        filename = f"nutriveda_plan_{submission_id}_{plan_id}_{safe_name}.docx"
        filepath = str(WORD_DIR / filename)
        doc.save(filepath)
        log.info(f"Word document saved: {filepath}")
        return filepath

    except Exception as e:
        log.error(f"Word document generation failed: {e}", exc_info=True)
        return None


# ── Admin Knowledge Source Report ─────────────────────────────────────────────

def generate_admin_doc(
    plan_id: int,
    submission_id: int,
    client_name: str,
    rag_sources: List[str],
    rag_chunks: List[dict],
    plan_generated_at: Optional[str] = None,
) -> Optional[str]:
    """
    Generate an admin-only Word document showing exactly which MHB knowledge base
    sources and text chunks were used to generate this diet plan.
    Helps verify that the local MHB study material folder is being consulted.
    Returns the saved file path, or None on failure.
    """
    try:
        doc = Document()
        _configure_document(doc)

        # ── Cover ──────────────────────────────────────────────────────────────
        # Admin header banner
        tbl = doc.add_table(rows=1, cols=1)
        _set_table_borders(tbl, color=str(C_SEC_NAVY), size=0)
        cell = tbl.cell(0, 0)
        _set_cell_bg(cell, C_SEC_NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("NutriVeda — Knowledge Source Report")
        r.bold = True
        r.font.size = Pt(20)
        r.font.color.rgb = C_WHITE

        p2 = cell.add_paragraph("FOR ADMIN USE ONLY — Nutrition Consultant Reference")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(16)
        r2 = p2.runs[0]
        r2.font.size = Pt(10)
        r2.italic = True
        r2.font.color.rgb = RGBColor(0xA0, 0xC4, 0xFF)

        _add_spacer(doc, 12)

        # Plan reference table
        ref_tbl = doc.add_table(rows=4, cols=2)
        _set_table_borders(ref_tbl, color=C_BORDER, size=4)
        ref_data = [
            ("Client", client_name),
            ("Plan Reference", f"NV-{submission_id:04d}  (Plan #{plan_id})"),
            ("Generated", plan_generated_at or datetime.now().strftime("%d %B %Y %H:%M")),
            ("MHB Sources Used", f"{len(rag_sources)} unique source files"),
        ]
        for row_idx, (label, value) in enumerate(ref_data):
            row = ref_tbl.rows[row_idx]
            _set_cell_bg(row.cells[0], C_HDR_BG)
            _cell_text(row.cells[0], label, bold=True, size=10,
                       color=C_WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell_bg(row.cells[1], C_LIGHT_GRN if row_idx % 2 == 0 else C_ROW_ALT)
            _cell_text(row.cells[1], value, size=10, color=C_TEXT)

        _add_spacer(doc, 16)

        # ── Section 1 — Source Files Summary ───────────────────────────────────
        _add_colored_heading(doc, "Section 1 — MHB Source Files Used", level=2,
                             banner_bg=C_SEC_NAVY)

        if not rag_sources:
            p = doc.add_paragraph()
            r = p.add_run(
                "⚠  No MHB knowledge base sources were retrieved for this plan.\n"
                "This means the RAG knowledge base was empty at generation time.\n"
                "Run build_rag.bat to index the MHB study material files, then regenerate."
            )
            r.font.size = Pt(11)
            r.font.color.rgb = C_SEC_AMBER
            r.bold = True
        else:
            for i, src in enumerate(rag_sources, 1):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(f"{i}.  {src}")
                r.font.size = Pt(10)
                r.font.color.rgb = C_TEXT

        _add_spacer(doc, 12)

        # ── Section 2 — Full Chunk Details ─────────────────────────────────────
        _add_colored_heading(doc, "Section 2 — Full Retrieved Text Chunks", level=2,
                             banner_bg=C_SEC_SLATE)

        if not rag_chunks:
            p = doc.add_paragraph()
            r = p.add_run(
                "No chunk data available. Chunk capture was added in a later version — "
                "regenerate the plan to capture full source data."
            )
            r.font.size = Pt(10)
            r.italic = True
            r.font.color.rgb = C_TEXT_GRAY
        else:
            for i, chunk in enumerate(rag_chunks, 1):
                source = chunk.get("source", "Unknown")
                topic = chunk.get("topic", "—")
                text = chunk.get("text", "")
                score = chunk.get("score", 0.0)
                relevance_pct = int(score * 100)

                # Chunk header
                ch_tbl = doc.add_table(rows=1, cols=1)
                _set_table_borders(ch_tbl, color=str(C_SEC_TEAL), size=4)
                ch_cell = ch_tbl.cell(0, 0)
                _set_cell_bg(ch_cell, C_SEC_TEAL)
                cp = ch_cell.paragraphs[0]
                cp.paragraph_format.space_before = Pt(4)
                cp.paragraph_format.space_after = Pt(4)
                cp.paragraph_format.left_indent = Pt(6)
                cr = cp.add_run(f"Chunk {i}   |   Source: {source}   |   Relevance: {relevance_pct}%")
                cr.bold = True
                cr.font.size = Pt(9)
                cr.font.color.rgb = C_WHITE

                _add_spacer(doc, 2)

                # Topic label
                tp = doc.add_paragraph()
                tp.paragraph_format.space_before = Pt(2)
                tp.paragraph_format.space_after = Pt(3)
                tp.paragraph_format.left_indent = Pt(8)
                tr = tp.add_run(f"Topic:  {topic}")
                tr.bold = True
                tr.italic = True
                tr.font.size = Pt(9.5)
                tr.font.color.rgb = C_SEC_TEAL

                # Text content box
                txt_tbl = doc.add_table(rows=1, cols=1)
                _set_table_borders(txt_tbl, color=C_BORDER, size=4)
                txt_cell = txt_tbl.cell(0, 0)
                _set_cell_bg(txt_cell, C_ROW_ALT)
                xp = txt_cell.paragraphs[0]
                xp.paragraph_format.space_before = Pt(5)
                xp.paragraph_format.space_after = Pt(5)
                xp.paragraph_format.left_indent = Pt(8)
                xp.paragraph_format.right_indent = Pt(8)
                xr = xp.add_run(text.strip() if text else "(no text content)")
                xr.font.size = Pt(9)
                xr.font.color.rgb = C_TEXT

                _add_spacer(doc, 10)

        # ── Section 3 — Verification Checklist ─────────────────────────────────
        _add_colored_heading(doc, "Section 3 — Verification Checklist", level=2,
                             banner_bg=C_SEC_GREEN)

        checks = [
            ("MHB sources retrieved", "Yes" if rag_chunks else "NO — RAG was empty"),
            ("Source files from local MHB folder", "Indexed via build_rag.bat" if rag_chunks else "Not indexed"),
            ("Number of chunks used", str(len(rag_chunks))),
            ("Unique source files", str(len(rag_sources))),
            ("Average relevance score", f"{sum(c.get('score',0) for c in rag_chunks)/max(len(rag_chunks),1)*100:.0f}%" if rag_chunks else "N/A"),
        ]
        check_tbl = doc.add_table(rows=len(checks), cols=2)
        _set_table_borders(check_tbl, color=C_BORDER, size=4)
        for row_idx, (label, val) in enumerate(checks):
            row = check_tbl.rows[row_idx]
            _set_cell_bg(row.cells[0], C_HDR_BG)
            _cell_text(row.cells[0], label, bold=True, size=10, color=C_WHITE)
            ok = "NO" not in val.upper() and "N/A" not in val.upper() and "NOT" not in val.upper()
            _set_cell_bg(row.cells[1], C_PALE_GRN if ok else RGBColor(0xFF, 0xE4, 0xE4))
            _cell_text(row.cells[1], val, size=10,
                       color=C_PRIMARY if ok else C_SEC_RED)

        _add_spacer(doc, 16)
        p_end = doc.add_paragraph()
        r_end = p_end.add_run(
            "End of Knowledge Source Report  —  NutriVeda Admin System  —  CONFIDENTIAL"
        )
        r_end.font.size = Pt(8)
        r_end.italic = True
        r_end.font.color.rgb = C_TEXT_GRAY
        p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save
        safe_name = re.sub(r"[^\w]", "_", client_name.lower())[:20]
        filename = f"nutriveda_admin_sources_{plan_id}_{safe_name}.docx"
        filepath = str(WORD_DIR / filename)
        doc.save(filepath)
        log.info(f"Admin source report saved: {filepath}")
        return filepath

    except Exception as e:
        log.error(f"Admin source report generation failed: {e}", exc_info=True)
        return None
